#!/usr/bin/env python3
"""One-time script to generate binary fixture files for CV reader tests.

Run once from the repo root:
    uv run python tests/fixtures/cvs/create_fixtures.py
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

FIXTURE_DIR = Path(__file__).parent


def make_minimal_pdf(text_lines: list[bytes]) -> bytes:
    """Build a valid minimal PDF with Helvetica text on one page."""
    content_parts = [b"BT", b"/F1 12 Tf", b"50 750 Td"]
    for line in text_lines:
        escaped = line.replace(b"(", b"\\(").replace(b")", b"\\)")
        content_parts.append(b"(" + escaped + b") Tj")
        content_parts.append(b"0 -20 Td")
    content_parts.append(b"ET")
    content = b"\n".join(content_parts)

    obj1 = b"<< /Type /Catalog /Pages 2 0 R >>"
    obj2 = b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>"
    obj3 = (
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
        b" /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
    )
    obj4 = b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream"
    obj5 = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

    buf = BytesIO()
    buf.write(b"%PDF-1.4\n")
    offsets: list[int] = []
    for obj in [obj1, obj2, obj3, obj4, obj5]:
        offsets.append(buf.tell())
        buf.write(str(len(offsets)).encode() + b" 0 obj\n" + obj + b"\nendobj\n\n")

    xref_offset = buf.tell()
    buf.write(b"xref\n")
    buf.write(b"0 " + str(len(offsets) + 1).encode() + b"\n")
    buf.write(b"0000000000 65535 f \n")
    for offset in offsets:
        buf.write(f"{offset:010d} 00000 n \n".encode())

    buf.write(
        b"trailer\n<< /Size "
        + str(len(offsets) + 1).encode()
        + b" /Root 1 0 R >>\n"
        + b"startxref\n"
        + str(xref_offset).encode()
        + b"\n%%EOF\n"
    )
    return buf.getvalue()


def create_pdf_fixture() -> None:
    pdf_bytes = make_minimal_pdf(
        [
            b"John Doe",
            b"john.doe@example.com | +1 555-0100",
            b"EDUCATION",
            b"BSc Computer Science, State University, 2018",
            b"EXPERIENCE",
            b"Software Engineer, Acme Corp (2018-2023)",
            b"SKILLS",
            b"Python, PostgreSQL, REST APIs, Docker, Git",
        ]
    )
    dest = FIXTURE_DIR / "sample.pdf"
    dest.write_bytes(pdf_bytes)
    print(f"Written {dest} ({len(pdf_bytes)} bytes)")


def create_docx_fixture() -> None:
    from docx import Document  # type: ignore[import-untyped]

    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("john.doe@example.com | +1 555-0100")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("BSc Computer Science, State University, 2018")
    doc.add_paragraph("EXPERIENCE")
    doc.add_paragraph("Software Engineer, Acme Corp (2018-2023)")

    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "PostgreSQL"

    dest = FIXTURE_DIR / "sample.docx"
    doc.save(str(dest))
    print(f"Written {dest}")


if __name__ == "__main__":
    create_pdf_fixture()
    create_docx_fixture()
    print("Done.")
