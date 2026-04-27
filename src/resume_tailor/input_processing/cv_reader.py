"""CV reader — extracts raw text from PDF, DOCX, and plain-text CV files."""

from __future__ import annotations

from pathlib import Path

__all__ = [
    "read_cv",
    "CVReaderError",
    "CVFileNotFoundError",
    "CVPermissionError",
    "CVUnsupportedFormatError",
    "CVEncodingError",
]

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt")


# ---------------------------------------------------------------------------
# Typed error hierarchy
# ---------------------------------------------------------------------------


class CVReaderError(Exception):
    """Base class for all CV reader errors."""


class CVFileNotFoundError(CVReaderError):
    """Raised when the CV file does not exist."""


class CVPermissionError(CVReaderError):
    """Raised when the CV file cannot be read due to a permission error."""


class CVUnsupportedFormatError(CVReaderError):
    """Raised when the file extension is not supported."""


class CVEncodingError(CVReaderError):
    """Raised when a plain-text file cannot be decoded as UTF-8."""


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def read_cv(path: str | Path) -> str:
    """Read a CV file and return its raw text content.

    Dispatches on the file extension (case-insensitive):
    - ``.pdf``  — extracted via pdfplumber
    - ``.docx`` — extracted via python-docx (paragraphs + table cells)
    - ``.txt``  — read as UTF-8

    Raises:
        CVFileNotFoundError: If the file does not exist.
        CVPermissionError: If the file cannot be read due to OS permissions.
        CVUnsupportedFormatError: If the extension is not one of the supported
            formats (.pdf, .docx, .txt).
        CVEncodingError: If a .txt file contains non-UTF-8 bytes.
    """
    path = Path(path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(SUPPORTED_EXTENSIONS)
        raise CVUnsupportedFormatError(
            f"Unsupported file format {path.suffix!r}. Supported formats: {supported}"
        )

    if not path.exists():
        raise CVFileNotFoundError(f"CV file not found: {path}")

    try:
        if ext == ".pdf":
            return _read_pdf(path)
        elif ext == ".docx":
            return _read_docx(path)
        else:
            return _read_txt(path)
    except PermissionError as exc:
        raise CVPermissionError(f"Permission denied reading CV file: {path}") from exc


# ---------------------------------------------------------------------------
# Format-specific readers
# ---------------------------------------------------------------------------


def _read_pdf(path: Path) -> str:
    """Extract text from a PDF using pdfplumber."""
    import pdfplumber

    with pdfplumber.open(path) as pdf:
        pages_text = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
    return "\n".join(pages_text)


def _read_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx.

    Includes paragraph text and table cell text.
    Headers, footers, and footnotes are out of scope.
    """
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        parts.append(para.text)

    return "\n".join(parts)


def _read_txt(path: Path) -> str:
    """Read a plain-text file as UTF-8."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise CVEncodingError(
            f"CV file is not valid UTF-8: {path}. Ensure the file is saved with UTF-8 encoding."
        ) from exc
