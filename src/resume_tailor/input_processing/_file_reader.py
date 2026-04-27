"""Shared file-reading helpers for PDF, DOCX, and plain-text files.

These functions contain the format-specific parsing logic used by both
``cv_reader`` and ``jd_reader``.  They raise standard built-in exceptions
(``PermissionError``, ``UnicodeDecodeError``) and do *not* raise typed
reader errors — that translation is the responsibility of each reader module.
"""

from __future__ import annotations

from pathlib import Path

__all__ = [
    "read_pdf",
    "read_docx",
    "read_txt",
    "SUPPORTED_EXTENSIONS",
]

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt")


def read_pdf(path: Path) -> str:
    """Extract text from a PDF using pdfplumber."""
    import pdfplumber

    with pdfplumber.open(path) as pdf:
        pages_text = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
    return "\n".join(pages_text)


def read_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx.

    Includes paragraph text and table cell text.
    Headers, footers, and footnotes are out of scope.
    """
    from docx import Document

    doc = Document(path)
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        parts.append(para.text)

    return "\n".join(parts)


def read_txt(path: Path) -> str:
    """Read a plain-text file as UTF-8.

    Re-raises ``UnicodeDecodeError`` as-is so that each calling reader module
    can catch it specifically and translate it into a domain-typed error.
    """
    return path.read_text(encoding="utf-8")
